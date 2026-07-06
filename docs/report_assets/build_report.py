# -*- coding: utf-8 -*-
"""
Build the SCMS project report (.docx) modelled on the reference
"College Forum App" report layout.

Output: docs/SCMS_Project_Report.docx
Run:    python docs/report_assets/build_report.py
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
ROOT = os.path.abspath(os.path.join(HERE, "..", ".."))
DIAG = os.path.join(HERE, "diagrams")
SHOT = os.path.join(ROOT, "docs", "screenshots", "ios-clean-redesign")
LOGO = os.path.join(HERE, "rvce_logo.png")   # RVCE letterhead (reused from reference)
OUT = os.path.join(ROOT, "docs", "SCMS_Project_Report.docx")

FONT = "Times New Roman"   # report body/heading font family

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
INK = RGBColor(0x1C, 0x25, 0x30)
GREY = RGBColor(0x5B, 0x6B, 0x7B)
BLACK = RGBColor(0x00, 0x00, 0x00)  # reference report uses plain black headings

# ---------------------------------------------------------------------------
# Project / team metadata  (EDIT the two placeholder USNs when known)
# ---------------------------------------------------------------------------
TITLE = "Smart Complaint Management System (SCMS)"
COURSE_NAME = "MOBILE APPLICATION DEVELOPMENT"
COURSE_CODE = "MCA221IA"
LAB = "SEE PROJECT BASED LABORATORY"
YEAR = "2025-2026"
GUIDES = "Prof. Chandrani Chakravorty & Prof. Prashanth K"
MEMBERS = [
    ("Pavan Naik", "1RV25MC066"),
    ("Prabhava Upadhyaya V", "1RV25MC068"),
    ("Prem Kamble", "1RV25MC074"),
    ("Pramath Hegde", "1RV25MC070"),
]

# ---------------------------------------------------------------------------
# Figures in document order.  (filename, caption, kind)  kind: diagram | shot
# ---------------------------------------------------------------------------
FIGS = [
    ("fig_architecture.png",   "Three-tier system architecture of SCMS",                       "diagram"),
    ("fig_feature_flow.png",   "Conceptual feature flow of SCMS",                              "diagram"),
    ("30-login-light.png",     "Login screen with Google Sign-In (restricted to rvce.edu.in)", "shot"),
    ("31-dashboard-light.png", "Student dashboard with live complaint data and status breakdown", "shot"),
    ("32-all-complaints-light.png", "Complaint list showing status chips, severity and SLA timers", "shot"),
    ("33-new-form-light.png",  "New-complaint form with subject, description and category chips", "shot"),
    ("34-grammar-suggestion-light.png", "AI grammar suggestion and AI category suggestion while composing", "shot"),
    ("35-grammar-applied-light.png",  "AI grammar correction applied and AI category suggestion (Plumbing, 100%)", "shot"),
    ("36-complaint-detail-light.png", "Complaint detail with details and status timeline",     "shot"),
    ("37-after-submit-light.png",     "Updated feed after submitting a complaint (Pending Review)", "shot"),
    ("38-stats-light.png",           "Statistics / analytics dashboard",                       "shot"),
    ("39-profile-light.png",         "Profile and grouped settings screen",                    "shot"),
    ("fig_navigation.png",     "Role-based navigation flow of SCMS",                           "diagram"),
    ("fig_ai_sequence.png",    "AI-assisted submission request sequence",                      "diagram"),
    ("fig_sr_workflow.png",    "Student Representative (SR) review workflow",                  "diagram"),
    ("fig_lifecycle.png",      "Complaint lifecycle and status flow",                          "diagram"),
]
FIGNUM = {name: i + 1 for i, (name, _, _) in enumerate(FIGS)}
FIGCAP = {name: cap for name, cap, _ in FIGS}


def fig_path(name):
    kind = next(k for n, c, k in FIGS if n == name)
    return os.path.join(DIAG if kind == "diagram" else SHOT, name)


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------
def set_cell_border_none(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tblPr.append(borders)


def add_field(paragraph, instr, placeholder):
    run = paragraph.add_run()
    r = run._r
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = placeholder
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    for el in (f1, it, f2, t, f3):
        r.append(el)


def add_running_header_footer(section, title, right_tab=Inches(6.5)):
    """Give a section a reference-style running header and footer.

    Header:  <title>                              <Month Year>
    Footer:  Department of MCA, RVCE                     Page <N>
    The date/right column is pushed flush-right via a right tab stop.
    """
    month_year = datetime.now().strftime("%B %Y")

    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    # ---- header ----
    hp = section.header.paragraphs[0]
    hp.text = ""
    hp.paragraph_format.tab_stops.add_tab_stop(right_tab, WD_TAB_ALIGNMENT.RIGHT)
    r1 = hp.add_run(title)
    r1.font.size = Pt(9)
    r1.font.color.rgb = GREY
    r2 = hp.add_run("\t" + month_year)
    r2.font.size = Pt(9)
    r2.font.color.rgb = GREY

    # ---- footer (with live PAGE field) ----
    fp = section.footer.paragraphs[0]
    fp.text = ""
    fp.paragraph_format.tab_stops.add_tab_stop(right_tab, WD_TAB_ALIGNMENT.RIGHT)
    fr = fp.add_run("Department of MCA, RVCE")
    fr.font.size = Pt(9)
    fr.font.color.rgb = GREY
    fp.add_run("\tPage ").font.size = Pt(9)
    add_field(fp, "PAGE", "1")


def para(doc, text="", size=11, align=None, bold=False, italic=False,
         color=None, space_after=6, space_before=0, line=1.15):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        if color:
            r.font.color.rgb = color
    return p


def bullets(doc, items, size=11):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        r.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(3)


def h1(doc, text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.color.rgb = BLACK
    return p


def h2(doc, text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.color.rgb = BLACK
    return p


def caption(doc, name):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(f"Figure {FIGNUM[name]}: {FIGCAP[name]}")
    r.font.size = Pt(9.5)
    r.font.italic = True
    r.font.color.rgb = GREY


def add_diagram(doc, name, width=6.3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(fig_path(name), width=Inches(width))
    caption(doc, name)


def add_shots_row(doc, names, width=2.35):
    """Place up to two phone screenshots side by side, captioned."""
    table = doc.add_table(rows=1, cols=len(names))
    table.alignment = 1  # center
    set_cell_border_none(table)
    for cell, name in zip(table.rows[0].cells, names):
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.add_run().add_picture(fig_path(name), width=Inches(width))
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f"Figure {FIGNUM[name]}: {FIGCAP[name]}")
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = GREY
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ===========================================================================
# Document setup
# ===========================================================================
doc = Document()

# base styles
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(11)
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.space_after = Pt(6)
# force the east-asian slot too so the theme's Calibri never leaks back in
normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)

for lvl, sz in ((1, 15), (2, 13)):
    st = doc.styles[f"Heading {lvl}"]
    st.font.name = FONT
    st.font.size = Pt(sz)
    st.font.bold = True
    st.font.color.rgb = BLACK

sec = doc.sections[0]
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)

# ---------------------------------------------------------------------------
# COVER PAGE
# ---------------------------------------------------------------------------
def add_logo(width=5.9):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    p.add_run().add_picture(LOGO, width=Inches(width))


def cover_line(text, size, bold=False, color=BLACK, after=4, before=0, italic=False):
    para(doc, text, size=size, align=WD_ALIGN_PARAGRAPH.CENTER, bold=bold,
         italic=italic, color=color, space_after=after, space_before=before, line=1.1)


def borderless_center_table(rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    t.alignment = 1  # center on page
    set_cell_border_none(t)
    return t


# ---- COVER (mirrors the reference report's layout) ----
add_logo()
cover_line(COURSE_NAME, 14, bold=True, after=0)          # MOBILE APPLICATION DEVELOPMENT
cover_line(COURSE_CODE, 14, bold=True, after=0)          # MCA221IA
cover_line(LAB, 14, bold=True, after=0)                  # SEE PROJECT BASED LABORATORY
cover_line(TITLE, 16, bold=True, before=6, after=10)
cover_line("submitted by", 11, after=6)

# student names + USNs in a borderless 2-column table
mt = borderless_center_table(len(MEMBERS), 2)
for i, (name, usn) in enumerate(MEMBERS):
    for j, val in enumerate((name, usn)):
        cp = mt.rows[i].cells[j].paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(val)
        r.font.size = Pt(11)

cover_line("under the guidance of", 11, before=8, after=4)
# guide in a borderless 1x1 table
gt = borderless_center_table(1, 1)
gp = gt.rows[0].cells[0].paragraphs[0]
gp.alignment = WD_ALIGN_PARAGRAPH.CENTER
gr = gp.add_run(GUIDES)
gr.font.size = Pt(11)
gr.font.bold = True

cover_line("Department of Master of Computer Applications", 11, bold=True, before=10, after=0)
cover_line("RV College of Engineering, Bengaluru – 560059", 11, bold=True, after=0)
cover_line(YEAR, 11, bold=True, after=0)

doc.add_page_break()

# ---------------------------------------------------------------------------
# CERTIFICATE (mirrors the reference report's layout)
# ---------------------------------------------------------------------------
add_logo()
para(doc, "CERTIFICATE", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True,
     color=BLACK, space_after=10)

names_join = ", ".join(f"{n} ({u})" for n, u in MEMBERS[:-1]) + \
    f" and {MEMBERS[-1][0]} ({MEMBERS[-1][1]})"
cert = (f"Certified that the project entitled “{TITLE}” on "
        f"{COURSE_NAME.title()} – {COURSE_CODE} has been carried out by "
        f"{names_join} who have successfully completed the project for the final "
        f"SEE Lab Examination, incorporating all concepts of the course conducted "
        f"by the Department of MCA, RV College of Engineering, Bengaluru.")
para(doc, cert, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10, line=1.5)

sig = doc.add_table(rows=1, cols=2)
set_cell_border_none(sig)
left = sig.rows[0].cells[0]
right = sig.rows[0].cells[1]
for cell, lines in ((left, ["Internal Guide", GUIDES, "Assistant Professor",
                            "Department of MCA", "RV College of Engineering"]),
                    (right, ["Head of the Department", "Dr. Jasmine K S",
                             "Associate Professor & Director", "Department of MCA,",
                             "RV College of Engineering"])):
    for i, ln in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        r = p.add_run(ln)
        r.font.size = Pt(11)
        r.font.bold = (i <= 1)

para(doc, "", space_after=20)
para(doc, "External Viva Examination", size=11, align=WD_ALIGN_PARAGRAPH.CENTER,
     bold=True, color=BLACK, space_after=10)
vp = doc.add_paragraph()
vr = vp.add_run("Name of Examiners\t\t\tSignature with Date")
vr.font.bold = True
vr.font.size = Pt(11)
for label in ("1.", "2."):
    lp = doc.add_paragraph()
    lp.add_run(label).font.size = Pt(11)

doc.add_page_break()

# ---------------------------------------------------------------------------
# TABLE OF CONTENTS
# ---------------------------------------------------------------------------
para(doc, "Table of Contents", size=16, align=WD_ALIGN_PARAGRAPH.CENTER,
     bold=True, color=BLACK, space_after=12)
tocp = doc.add_paragraph()
add_field(tocp, 'TOC \\o "1-2" \\h \\z \\u',
          "Right-click here and choose “Update Field” to build the table of contents.")
doc.add_page_break()

# ---------------------------------------------------------------------------
# LIST OF FIGURES  (static, generated from FIGS)
# ---------------------------------------------------------------------------
para(doc, "List of Figures", size=16, align=WD_ALIGN_PARAGRAPH.CENTER,
     bold=True, color=BLACK, space_after=12)
for name, cap, _ in FIGS:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"Figure {FIGNUM[name]}:  ")
    r.font.bold = True
    r.font.size = Pt(10.5)
    r2 = p.add_run(cap)
    r2.font.size = Pt(10.5)

# Start a new section for the report body so the front matter (cover, certificate,
# TOC, list of figures) stays free of the running header/footer, while every body
# page carries the reference-style header + footer with live page numbers.
body_section = doc.add_section(WD_SECTION.NEW_PAGE)
add_running_header_footer(body_section, TITLE)

# ===========================================================================
# 1. INTRODUCTION
# ===========================================================================
h1(doc, "1. Introduction")
para(doc,
     "The Smart Complaint Management System (SCMS) is a mobile-first platform "
     "developed to digitise and streamline the way complaints are raised, routed, "
     "tracked and resolved within an educational campus. It replaces informal, "
     "untracked channels — verbal requests, paper registers, scattered chat "
     "groups — with a single, accountable workflow in which every complaint has "
     "an owner, a status and a service-level deadline. Within the app, students "
     "raise complaints about issues such as plumbing, electrical faults, "
     "housekeeping, civil maintenance and IT, attach geotagged photographic "
     "evidence, and follow the progress of each complaint from submission to "
     "resolution.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
para(doc,
     "SCMS is built as three cooperating services. A Flutter mobile application "
     "provides the user interface for all four roles; a Node.js / Express backend "
     "exposes the REST API, enforces authentication and authorisation, and runs "
     "scheduled background jobs; and a Python (FastAPI) AI microservice augments "
     "the workflow with Google Gemini for grammar correction, automatic "
     "categorisation and duplicate-complaint detection. The mobile app never "
     "contacts the AI service directly — all AI requests are proxied through the "
     "Node.js backend, keeping the microservice internal and best-effort. Data is "
     "persisted in PostgreSQL, extended with the pgvector extension so that complaint "
     "descriptions can be stored as embeddings and compared for similarity.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
para(doc,
     "The system implements role-based access control with four roles — Student, "
     "Staff, Student Representative (SR) and Admin. Students submit and rate "
     "complaints; SRs review submissions and assign them to the appropriate staff; "
     "staff act on and resolve complaints; and admins manage users, departments and "
     "analytics. Authentication is through Google OAuth 2.0 restricted to the "
     "institutional domain rvce.edu.in, so only verified campus accounts can sign "
     "in. Figure 1 shows how these services are arranged.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_diagram(doc, "fig_architecture.png")

h2(doc, "1.2 Overview of Problem Domain")
para(doc,
     "Campuses generate a continuous stream of maintenance and service issues — a "
     "leaking tap, a broken projector, a power outage in a lab, an uncleaned "
     "corridor. Traditionally these are reported informally and handled without any "
     "systematic record. This creates several recurring problems: complaints are "
     "lost or forgotten, there is no accountability for who should act, students "
     "have no visibility into progress, and the same issue is often reported many "
     "times by different people.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
para(doc,
     "The absence of structure also makes prioritisation impossible. Without a "
     "severity level or a service-level agreement (SLA), an urgent electrical hazard "
     "receives the same attention as a minor cosmetic issue. Manual routing means "
     "complaints frequently reach the wrong department, adding delay. And because "
     "there is no consolidated data, management cannot see trends, recurring "
     "problem areas or staff workload.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
para(doc,
     "SCMS addresses this problem domain by providing a centralised, role-based "
     "workflow with enforced accountability. Every complaint is authenticated to a "
     "real campus user, routed through a review gate, assigned to a responsible "
     "member of staff, tracked against an SLA deadline, and checked for duplicates "
     "before it enters the queue. AI assistance improves the quality and "
     "consistency of submissions, while analytics give administrators a live view "
     "of the campus’s complaint landscape.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# ===========================================================================
# 2. KEY FEATURES AND ANALYSIS
# ===========================================================================
h1(doc, "2. Key Features and Analysis")
para(doc,
     "SCMS combines secure authentication, structured complaint handling, "
     "AI augmentation and analytics into a single platform. The major features are "
     "described below along with a short analysis of the value each one provides. "
     "Figure 2 summarises how these features relate to the platform.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "2.1 Role-Based Access Control")
para(doc, "The system supports four distinct roles, each backed by a JWT claim "
     "(ROLE_USER, ROLE_STAFF, ROLE_SR, ROLE_ADMIN) and guarded on the server:")
bullets(doc, [
    "Student (User): raises complaints, tracks their status and rates the resolution.",
    "Staff: sees complaints assigned to them and updates their status through to resolution.",
    "Student Representative (SR): reviews newly submitted complaints and assigns them to staff.",
    "Admin: manages users, departments, categories and zones, and views analytics.",
])
para(doc,
     "Route-level guards ensure that each role can reach only the actions it is "
     "permitted to perform, preventing, for example, a student from changing a "
     "complaint’s status or an unauthenticated request from reaching protected "
     "endpoints.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "2.2 Secure Authentication with Google OAuth and JWT")
para(doc,
     "Sign-in is exclusively through Google OAuth 2.0, restricted to the "
     "rvce.edu.in hosted domain — there is no email/password login. The Google ID "
     "token returned to the app is verified on the backend, which additionally "
     "enforces an email-domain allowlist before issuing its own short-lived access "
     "token and a refresh token. The mobile app stores these tokens in secure "
     "storage and transparently refreshes the access token when it expires, so "
     "sessions remain both secure and seamless.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "2.3 Complaint Submission with Geotagged, Watermarked Photos")
para(doc,
     "When raising a complaint, a student provides a subject, a description, a "
     "severity level and an optional location, and may attach up to three photos as "
     "evidence. Each photo is stamped with the capture GPS coordinates and a "
     "date-time watermark before upload, which discourages the reuse of old or "
     "unrelated images and gives staff reliable context about where and when the "
     "issue was observed.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "2.4 AI Grammar Correction")
para(doc,
     "As the student types the description, the app debounces the input and asks "
     "the backend to run a grammar and clarity pass through the AI service (Google "
     "Gemini). The corrected text is offered back as a suggestion the student can "
     "accept with a single tap. This raises the quality and readability of "
     "complaints without forcing the student to rewrite them, which in turn helps "
     "staff understand issues faster.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "2.5 AI Category Suggestion")
para(doc,
     "In parallel with the grammar check, the AI service classifies the complaint "
     "text into one of the campus categories (for example Plumbing, Electrical, "
     "Housekeeping or Civil) and returns a suggested category with a confidence "
     "score and a short justification. The student can accept the suggestion or "
     "choose a different category. Accurate categorisation at source improves "
     "routing and reduces the chance of a complaint reaching the wrong department.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "2.6 Duplicate-Complaint Detection")
para(doc,
     "After a complaint row is created, the backend asks the AI service to compute "
     "a 768-dimensional embedding of its description and store it in the pgvector "
     "column. New complaints are compared against existing embeddings using cosine "
     "similarity; when a submission is highly similar to an earlier one it can be "
     "flagged as a probable duplicate. This prevents the queue from filling with "
     "repeated reports of the same underlying issue.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
para(doc,
     "The AI service is designed to fail safe: if Gemini or the database is "
     "unavailable, each endpoint returns a safe default rather than an error, and "
     "the complaint is still created normally. AI is treated as best-effort "
     "augmentation, never a hard dependency of the core workflow.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "2.7 SR Review and Assignment Workflow")
para(doc,
     "Newly submitted complaints enter a PENDING_SR_REVIEW state and appear in the "
     "Student Representative’s queue. The SR verifies the complaint, and either "
     "returns it to the student for clarification or approves it and assigns it to "
     "the responsible staff member or department. This human review gate keeps the "
     "queue clean and ensures complaints are correctly routed before staff are "
     "asked to act on them.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "2.8 SLA Tracking and Automatic Escalation")
para(doc,
     "Every complaint carries a service-level deadline based on its severity. "
     "Scheduled background jobs run periodically on the backend: one marks "
     "complaints whose deadline has passed as SLA_BREACHED so that overdue issues "
     "are visible and can be escalated, and another auto-approves complaints that "
     "have been stuck in SR review beyond a threshold so that the pipeline never "
     "stalls. SLA tracking turns vague promises of “we’ll look into it” into "
     "measurable, enforced commitments.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "2.9 Real-Time Push Notifications")
para(doc,
     "SCMS uses Firebase Cloud Messaging to deliver push notifications for "
     "important events — a complaint being assigned, its status changing, or a "
     "resolution being recorded — even when the app is in the background. In-app "
     "banners and deep links take the user straight to the relevant complaint. This "
     "keeps every party informed without anyone having to poll the app for updates.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "2.10 Offline Draft Support")
para(doc,
     "Because campus connectivity is not always reliable, the app can save a "
     "complaint as a local draft using an on-device store. If the network is "
     "unavailable at submission time, the repository falls back to the saved draft, "
     "and the complaint can be completed once connectivity returns. This ensures "
     "that a student is never blocked from at least capturing an issue.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "2.11 Ratings and Analytics")
para(doc,
     "After a complaint is resolved, the student can rate the resolution and leave "
     "a comment, providing a feedback signal on service quality. Administrators see "
     "aggregated analytics — complaint volumes, statuses, category distribution and "
     "SLA performance — giving management the data needed to identify recurring "
     "problems and allocate resources.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_diagram(doc, "fig_feature_flow.png")

# ===========================================================================
# 3. TECHNOLOGY STACK
# ===========================================================================
h1(doc, "3. Technology Stack")
para(doc,
     "SCMS is built on a modern, layered stack spanning a cross-platform mobile "
     "client, a Node.js API, a Python AI microservice and a vector-capable "
     "relational database. The technologies were chosen for security, "
     "responsiveness and the ability to integrate AI without compromising the "
     "reliability of the core workflow.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "3.1 Mobile Application Technologies")
para(doc, "Flutter (Dart): a single codebase compiled to native Android (and iOS) "
     "builds, providing the entire user interface and a consistent, responsive "
     "experience with light and dark themes.")
bullets(doc, [
    "BLoC / Cubit — predictable state management, one bloc per feature (auth, complaint, submission, SR review, analytics).",
    "Dio — HTTP client with interceptors that attach the JWT and transparently unwrap the backend response envelope.",
    "go_router — declarative routing with role-based redirects so each role sees only its permitted screens.",
    "Hive — lightweight on-device storage for offline complaint drafts.",
    "flutter_secure_storage — encrypted storage for access and refresh tokens.",
])

h2(doc, "3.2 Backend Technologies")
para(doc, "Node.js with the Express framework provides the REST API on port 3000.")
bullets(doc, [
    "Express — routing, middleware, static media serving, and a uniform success/error response envelope.",
    "Prisma ORM — type-safe database access and schema migrations against PostgreSQL.",
    "node-cron — scheduled jobs for SLA-breach marking and SR auto-approval.",
    "Helmet, CORS and Morgan — security headers, cross-origin control and request logging.",
])

h2(doc, "3.3 AI Service Technologies")
para(doc, "A Python FastAPI microservice on port 8000 encapsulates all AI logic.")
bullets(doc, [
    "FastAPI — fast, typed HTTP endpoints for grammar-check, categorize, embed and check-duplicate.",
    "google-genai (Google Gemini) — gemini-2.5-flash for text tasks and gemini-embedding-004 for 768-dimensional embeddings.",
    "psycopg2 with pgvector — pooled database access and cosine-similarity search for duplicate detection.",
])

h2(doc, "3.4 Database Technology")
para(doc,
     "PostgreSQL is the primary datastore, extended with the pgvector extension. "
     "Relational tables model users, complaints, media items, status-update "
     "timelines, departments, categories, zones, tags and refresh tokens, while the "
     "vector column stores complaint-description embeddings for similarity search. "
     "Running Postgres and pgvector together in one engine keeps transactional data "
     "and vector search consistent and simple to operate.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "3.5 Authentication and Security Technologies")
bullets(doc, [
    "Google OAuth 2.0 (google-auth-library) — identity restricted to the rvce.edu.in domain.",
    "JSON Web Tokens — stateless access tokens and rotating refresh tokens with role claims.",
    "Domain allowlist — a second server-side check of the signed-in email domain.",
    "Helmet and HTTPS — secure headers and encrypted transport.",
])

h2(doc, "3.6 Development and Deployment Tools")
bullets(doc, [
    "Docker Compose — runs the pgvector-enabled PostgreSQL container.",
    "Firebase (firebase-admin) — push notification delivery via Cloud Messaging.",
    "nodemon — backend hot-reload during development.",
    "Prisma Studio — GUI inspection of the database.",
    "Git — version control with a documented per-member work division.",
])

# ===========================================================================
# 4. REQUIREMENTS & SPECIFICATIONS
# ===========================================================================
h1(doc, "4. Requirements & Specifications")

h2(doc, "4.1 Functional Requirements")
bullets(doc, [
    "Users must authenticate through Google OAuth restricted to the rvce.edu.in domain.",
    "Students can create a complaint with a subject, description, severity, optional location and up to three photos.",
    "The system watermarks each attached photo with GPS coordinates and a timestamp.",
    "The app offers AI grammar correction and an AI-suggested category while composing a complaint.",
    "The system detects probable duplicate complaints using description embeddings.",
    "Newly submitted complaints enter SR review before assignment.",
    "SRs can approve and assign complaints to staff, or return them to the student.",
    "Staff can update the status of assigned complaints through to resolution.",
    "The system tracks an SLA deadline per complaint and marks breaches automatically.",
    "The system sends push notifications on assignment, status change and resolution.",
    "Students can rate a resolved complaint and leave a comment.",
    "Admins can manage users, departments, categories and zones, and view analytics.",
])

h2(doc, "4.2 Non-Functional Requirements")
bullets(doc, [
    "Security — all protected endpoints require a valid JWT; roles are enforced server-side.",
    "Reliability — AI features fail safe; core complaint handling continues if the AI service is down.",
    "Performance — typing-time AI calls are debounced to avoid excessive requests.",
    "Usability — mobile-first, responsive UI with light and dark themes.",
    "Availability — offline drafts allow complaints to be captured without connectivity.",
    "Maintainability — layered architecture with clear module and file-ownership boundaries.",
    "Scalability — stateless JWT auth and a separable AI microservice.",
])

h2(doc, "4.3 System Specifications")
para(doc, "Software:")
bullets(doc, [
    "Mobile: Flutter SDK, Android (API level per device); Dart.",
    "Backend: Node.js with Express; Prisma; runs on port 3000.",
    "AI service: Python 3 with FastAPI and Uvicorn; runs on port 8000.",
    "Database: PostgreSQL 16 with the pgvector extension (via Docker Compose).",
    "External services: Google OAuth, Google Gemini API, Firebase Cloud Messaging.",
])
para(doc, "Hardware (development / demo):")
bullets(doc, [
    "An Android smartphone or emulator for the client.",
    "A development machine hosting the backend, AI service and database.",
    "Network connectivity between the phone and the backend for live use.",
])

# ===========================================================================
# 5. UI SCREENSHOTS
# ===========================================================================
h1(doc, "5. User Interface Screens")
para(doc,
     "The SCMS interface follows a clean, iOS-inspired visual language with grouped "
     "sections, clear status chips and full light/dark theme support. The following "
     "screens illustrate the main student-facing flows and the AI assistance. "
     "Screens for the Staff, SR and Admin roles share the same components and are "
     "described through the workflow diagrams in Section 6.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "5.1 Authentication and Dashboard")
para(doc,
     "Users sign in with Google, restricted to institutional accounts, and on "
     "success are redirected to the dashboard appropriate to their role. The "
     "student dashboard summarises complaint counts, highlights items that need "
     "attention and shows a live status breakdown.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_shots_row(doc, ["30-login-light.png", "31-dashboard-light.png"])

h2(doc, "5.2 Complaint List and Submission")
para(doc,
     "The complaint list shows every complaint with a status chip, severity, "
     "category and SLA timer, and can be searched and filtered. The new-complaint "
     "form collects the subject, description, location, category and severity, and "
     "allows up to three photos to be attached.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_shots_row(doc, ["32-all-complaints-light.png", "33-new-form-light.png"])

h2(doc, "5.3 AI Assistance — Grammar and Category")
para(doc,
     "As the description is typed, the AI grammar suggestion and the AI category "
     "suggestion appear. Figure 8 shows the corrected description applied together "
     "with the suggested category (Plumbing) at 100% confidence and a short "
     "justification.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_shots_row(doc, ["34-grammar-suggestion-light.png", "35-grammar-applied-light.png"])

h2(doc, "5.4 Complaint Detail and Analytics")
para(doc,
     "Each complaint has a detail view with its full metadata and status timeline. "
     "After submission the feed updates to reflect the new complaint, the statistics "
     "screen presents aggregated analytics, and the profile screen groups account "
     "and preference settings.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_shots_row(doc, ["36-complaint-detail-light.png", "37-after-submit-light.png"])
add_shots_row(doc, ["38-stats-light.png", "39-profile-light.png"])

h2(doc, "5.5 Navigation Flow")
para(doc,
     "Navigation is driven by the authenticated role. After sign-in, go_router "
     "redirects the user to the correct role home and exposes only the routes that "
     "role is allowed to use, as shown in Figure 13.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_diagram(doc, "fig_navigation.png")

# ===========================================================================
# 6. IMPLEMENTATION DETAILS
# ===========================================================================
h1(doc, "6. Implementation Details")
para(doc,
     "The system is organised into cohesive modules across the three services. Each "
     "module has a single responsibility and communicates through well-defined HTTP "
     "interfaces and the shared response envelope.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "6.1 Authentication Module")
para(doc,
     "The Flutter app performs Google Sign-In and sends the resulting ID token to "
     "the backend’s /api/auth/google endpoint. The backend verifies the token with "
     "Google, checks the email domain against the allowlist, provisions or looks up "
     "the user, and issues its own access and refresh JWTs carrying the user’s "
     "role. A Dio interceptor on the client attaches the access token to every "
     "request and refreshes it automatically on expiry.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "6.2 Complaint Submission Module")
para(doc,
     "Complaint creation is a multipart request carrying the subject, description, "
     "severity, optional location, JSON-encoded tags and the media files. Photos are "
     "watermarked with GPS and timestamp on the device before upload and stored by "
     "the backend’s storage service. If a category is chosen, the backend resolves "
     "the responsible department from the category; complaints are created in the "
     "SUBMITTED / PENDING_SR_REVIEW state.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "6.3 AI Proxy Module")
para(doc,
     "All AI functionality is reached through a single backend proxy that forwards "
     "requests to the Python service’s /grammar-check, /categorize, /embed and "
     "/check-duplicate endpoints. Grammar and categorisation run while the user "
     "types (debounced by 800 ms); embedding and duplicate checks run after the "
     "complaint row exists. Every AI endpoint fails safe, returning a usable "
     "default if Gemini or the database is unavailable. Figure 14 shows the request "
     "sequence.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_diagram(doc, "fig_ai_sequence.png")

h2(doc, "6.4 SR Review and Assignment Module")
para(doc,
     "The SR module lists complaints awaiting review, and lets the SR approve and "
     "assign a complaint to staff or return it to the student. On assignment the "
     "complaint moves to ASSIGNED and the assigned staff member is notified. "
     "Figure 15 shows the end-to-end review workflow across the student, AI service, "
     "SR and staff.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_diagram(doc, "fig_sr_workflow.png")

h2(doc, "6.5 Status and SLA Module")
para(doc,
     "Staff advance a complaint through IN_PROGRESS to RESOLVED, and each change is "
     "recorded on the complaint’s status timeline. Two node-cron jobs support the "
     "SLA policy: one marks overdue complaints as SLA_BREACHED, and another "
     "auto-approves complaints left too long in SR review. Figure 16 shows the full "
     "status lifecycle.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_diagram(doc, "fig_lifecycle.png")

h2(doc, "6.6 Notification Module")
para(doc,
     "The backend sends push notifications through Firebase Cloud Messaging when a "
     "complaint is assigned, changes status or is resolved. On the client a "
     "notification service surfaces these as system notifications and in-app "
     "banners, and deep-links the user to the relevant complaint.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "6.7 Media and Watermark Module")
para(doc,
     "A dedicated watermark service stamps GPS coordinates and the capture "
     "date-time onto each photo using a custom painter before the image is uploaded. "
     "The backend stores media locally and serves it statically, and an enrichment "
     "helper attaches photo URLs and denormalised display fields to complaint "
     "responses.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "6.8 Offline Draft Module")
para(doc,
     "The complaint repository combines the remote data source with a local Hive "
     "store. When the network is unavailable it falls back to a saved draft, "
     "allowing a complaint to be captured offline and completed once connectivity "
     "is restored.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "6.9 API and Response-Envelope Layer")
para(doc,
     "Every backend response is wrapped in a uniform envelope — a success flag "
     "with a data payload, or an error object with a message and code. A client-side "
     "interceptor strips this envelope so that data sources receive the raw payload "
     "directly. This single convention keeps error handling consistent across the "
     "whole API surface.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "6.10 Overall System Workflow")
para(doc,
     "Bringing the modules together, a typical complaint flows as follows: the "
     "student authenticates with Google; composes a complaint assisted by AI grammar "
     "and category suggestions; submits it with watermarked photos; the backend "
     "stores it, embeds it and checks for duplicates; the SR reviews and assigns it; "
     "staff are notified and work the complaint to resolution while the SLA job "
     "watches the deadline; and finally the student is notified and rates the "
     "outcome. The architecture in Figure 1 and the workflow in Figures 15 and 16 "
     "together describe this end-to-end path.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# ===========================================================================
# 7. DEMONSTRATION
# ===========================================================================
h1(doc, "7. Demonstration")

h2(doc, "7.1 Demonstration Description")
para(doc,
     "The demonstration runs the complete stack — the Flutter app on an Android "
     "device, the Node.js backend, the Python AI service and the PostgreSQL / "
     "pgvector database — and walks a single complaint through the full lifecycle "
     "across all four roles.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "7.2 Demonstration Flow")
bullets(doc, [
    "Sign in as a student with a Google rvce.edu.in account.",
    "Start a new complaint; observe AI grammar correction and the AI-suggested category appear while typing.",
    "Attach a photo and see the GPS + date-time watermark applied.",
    "Submit the complaint; the backend stores it, embeds it and runs a duplicate check.",
    "Sign in as the SR; review the pending complaint and assign it to a staff member.",
    "Sign in as staff; move the complaint through In Progress to Resolved and observe the status timeline update.",
    "Receive push notifications on status changes; as the student, rate the resolved complaint.",
    "Sign in as admin; view the analytics dashboard reflecting the new activity.",
])

h2(doc, "7.3 Demonstration Notes")
para(doc,
     "The screenshots in Section 5 are captured from live runs of the application "
     "against the running backend and AI service, including the AI grammar "
     "correction and category suggestion shown in Figure 8. Because the AI service "
     "is best-effort, the demonstration also shows the system continuing to function "
     "normally when an AI response is temporarily unavailable.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# ===========================================================================
# 8. CONCLUSION
# ===========================================================================
h1(doc, "8. Conclusion")

h2(doc, "8.1 Contributions of the Project")
para(doc,
     "SCMS delivers a complete, accountable complaint-management workflow for a "
     "campus, replacing informal reporting with an authenticated, role-based, "
     "SLA-tracked system. Its distinctive contribution is the pragmatic integration "
     "of AI — grammar correction, automatic categorisation and embedding-based "
     "duplicate detection — as a best-effort augmentation layer that improves "
     "submission quality and routing without ever becoming a single point of "
     "failure for the core service.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h2(doc, "8.2 Academic Learning Outcomes")
bullets(doc, [
    "Designing and building a cross-platform mobile app in Flutter with clean, layered architecture and BLoC state management.",
    "Building a secure REST API in Node.js/Express with JWT auth, Prisma and a consistent response envelope.",
    "Integrating a Python FastAPI microservice and a third-party LLM (Google Gemini) as a fail-safe augmentation.",
    "Applying vector search (pgvector) to a real problem — duplicate detection via embeddings.",
    "Implementing OAuth 2.0 with domain restriction, and push notifications with Firebase Cloud Messaging.",
    "Coordinating a four-member team with clear file-ownership boundaries and shared project documentation.",
])

h2(doc, "8.3 Future Scope")
bullets(doc, [
    "A web/admin console mirroring the mobile analytics for larger-screen management.",
    "Richer analytics — SLA compliance trends, category heat-maps and staff workload balancing.",
    "Smarter routing that learns the best department/staff from historical resolutions.",
    "Configurable, category-specific SLA policies and escalation chains.",
    "In-app chat between students and staff on a complaint thread.",
    "Multi-institution support with per-tenant domains and departments.",
])

# ===========================================================================
# 9. REFERENCES
# ===========================================================================
h1(doc, "9. References")
refs = [
    "Flutter Documentation — https://docs.flutter.dev",
    "Dart Language — https://dart.dev",
    "Node.js Documentation — https://nodejs.org/en/docs",
    "Express Framework — https://expressjs.com",
    "Prisma ORM — https://www.prisma.io/docs",
    "PostgreSQL — https://www.postgresql.org/docs",
    "pgvector Extension — https://github.com/pgvector/pgvector",
    "FastAPI — https://fastapi.tiangolo.com",
    "Google Gemini API (google-genai) — https://ai.google.dev",
    "Google Identity / OAuth 2.0 — https://developers.google.com/identity",
    "Firebase Cloud Messaging — https://firebase.google.com/docs/cloud-messaging",
    "JSON Web Tokens — https://jwt.io",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"[{i}]  {r}")
    run.font.size = Pt(10.5)

doc.save(OUT)
print("Saved:", OUT)
